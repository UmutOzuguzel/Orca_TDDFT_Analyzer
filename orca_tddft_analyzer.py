from docx import Document
from docx.shared import Pt, RGBColor
import os
import re
import pandas as pd
import numpy as np

#change below values for the gaussian
orca_outputfile = 'TDDFT.out'
x_start = 10000
x_end = 35000
x_resolution = 125
mode = "g" # "g" for gaussian, "l" for lorentzian curves


def extract_and_correct_tddft(input_file_path):
    # Load the TDDFT.out file
    with open(input_file_path, 'r') as file:
        lines = [line.strip() for line in file.readlines()]

    # Find the start of the relevant data for ABSORPTION SPECTRUM VIA TRANSITION ELECTRIC DIPOLE MOMENTS
    start_line = [i for i, line in enumerate(lines) if 'ABSORPTION SPECTRUM VIA TRANSITION ELECTRIC DIPOLE MOMENTS' in line][0] + 5

    # Extract data using regex
    pattern = re.compile(r'(\d+)\s+(-?\d+\.\d+)\s+(-?\d+\.\d+)\s+(-?\d+\.\d+)\s+(-?\d+\.\d+)\s+(-?\d+\.\d+)\s+(-?\d+\.\d+)\s+(-?\d+\.\d+)')
    extracted_data = []
    for line in lines[start_line:]:
        match = pattern.search(line)
        if match:
            extracted_data.append(match.groups())
        else:
            break  # stop when we reach a line that doesn't match the pattern

    # Convert to DataFrame
    columns = ['State', 'Energy (cm-1)', 'Wavelength', 'fosc', 'T2', 'TX', 'TY', 'TZ']
    df = pd.DataFrame(extracted_data, columns=columns)

    # Convert the 'Wavelength' column to eV
    df['Wavelength'] = df['Wavelength'].astype(float)
    df['eV'] = 1239.84193 / df['Wavelength']

    # Update the 'eV / nm' column to the desired format
    df['eV / nm'] = df['eV'].round(3).astype(str) + " (" + df['Wavelength'].round(2).astype(str) + ")"

    # Reorder the columns
    df = df[['State', 'Energy (cm-1)', 'eV / nm', 'fosc', 'T2', 'TX', 'TY', 'TZ']]

    # Extract Cartesian coordinates
    start_line = [i for i, line in enumerate(lines) if 'CARTESIAN COORDINATES (ANGSTROEM)' in line][0] + 2
    end_line = [i for i, line in enumerate(lines[start_line:], start=start_line) if line == ''][0]

    cartesian_coordinates = lines[start_line:end_line]

    # Extract the FINAL SINGLE POINT ENERGY
    energy_line = [line for line in lines if "FINAL SINGLE POINT ENERGY" in line][0]
    final_single_point_energy = energy_line.split()[-1]

    return df, cartesian_coordinates, final_single_point_energy


def write_to_docx(df, cartesian_coordinates, energy, output_file_path):
    doc = Document()

    # Write the title for the absorption data
    title = doc.add_paragraph()
    run = title.add_run("Electronic Absorption")
    run.font.name = 'Times New Roman'  # set font to Times New Roman
    run.font.size = Pt(14)  # set font size to 14
    run.font.bold = True  # make the title bold
    run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)  # set font color to a shade of blue

    # Create a table for the DataFrame
    df_subset = df[['State', 'Energy (cm-1)', 'eV / nm', 'fosc']]  # subset DataFrame
    table = doc.add_table(rows=1, cols=len(df_subset.columns))

    # Write the DataFrame to the table
    for i, column in enumerate(df_subset.columns):
        table.cell(0, i).text = column
    for i, row in df_subset.iterrows():
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = str(value)

    # Write the energy
    energy_paragraph = doc.add_paragraph(f"\nFINAL SINGLE POINT ENERGY: {energy}")
    for run in energy_paragraph.runs:
        run.font.name = 'Arial'  # set font to Arial
        run.font.size = Pt(12)  # set font size to 12

    # Write the Cartesian coordinates
    doc.add_paragraph("\nCARTESIAN COORDINATES (ANGSTROEM):")
    for line in cartesian_coordinates:
        coordinates_paragraph = doc.add_paragraph(line)
        for run in coordinates_paragraph.runs:
            run.font.name = 'Arial'  # set font to Arial
            run.font.size = Pt(11)  # set font size to 11

    # Save the document
    doc.save(output_file_path)

def main():
    # Define the path to the TDDFT.out file
    input_file_path = os.path.join(os.getcwd(), orca_outputfile)

    # Extract the relevant data
    df_result, cartesian_coordinates, final_single_point_energy = extract_and_correct_tddft(input_file_path)

    # Write the data to a CSV file
    df_result.to_csv(os.path.join(os.getcwd(), 'absorption_spectrum.csv'), index=False)

    # Write the data to a Word file
    write_to_docx(df_result, cartesian_coordinates, final_single_point_energy, os.path.join(os.getcwd(), 'tddft_data.docx'))

# Call the main function
if __name__ == "__main__":
    main()

import pandas as pd
import os

# Define the paths for the CSV and the text file
csv_path = os.path.join(os.getcwd(), 'absorption_spectrum.csv')
txt_path = os.path.join(os.getcwd(), 'Absorption_Input.txt')

print(f"Reading from: {csv_path}")
print(f"Writing to: {txt_path}")

try:
    # Read the data from the CSV
    df = pd.read_csv(csv_path)
except FileNotFoundError:
    print("CSV file not found.")
    exit()

# Open the text file for writing
try:
    with open(txt_path, 'w') as f:
        # Loop through each row and write formatted data to the text file
        for _, row in df.iterrows():
            energy = row['Energy (cm-1)']
            fosc = row['fosc']
            f.write(f"{energy} {fosc} 1500\n")
    print(f"Data written to {txt_path}")
except Exception as e:
    print(f"An error occurred: {e}")

# Verify that the data has been written correctly by reading the text file
try:
    with open(txt_path, 'r') as f:
        print(f.read())
except Exception as e:
    print(f"An error occurred while reading the text file: {e}")


import numpy as np
import matplotlib.pyplot as plt

def gaussian_fit(interpolated_x, peak_x, peak_y, fwhm):
    """Compute the Gaussian distribution based on given parameters."""
    std_dev = fwhm / 2.35482
    gaussian = peak_y * (1 / (std_dev * np.sqrt(2*np.pi))) * np.exp(-((interpolated_x - peak_x)**2)/(2*(std_dev**2)))
    gaussian = gaussian / (gaussian.max() + 1e-12) * peak_y
    return gaussian

def lorentzian_fit(interpolated_x, peak_x, peak_y, fwhm):
    """Compute the Lorentzian distribution based on given parameters."""
    hwhm = fwhm / 2.0
    lorentzian = peak_y * (1/(np.pi*hwhm)) * (hwhm**2) / ((interpolated_x-peak_x)**2 + hwhm**2)
    lorentzian = lorentzian / (lorentzian.max() + 1e-12) * peak_y
    return lorentzian

def read_data(filename):
    """Read data from a text file and return as numpy arrays."""
    x, y, fwhm = [], [], []
    with open(filename, 'r') as infile:
        for line in infile:
            x_val, y_val, fwhm_val = line.replace('\r', '').replace('\n','').split(' ')
            x.append(float(x_val))
            y.append(float(y_val))
            fwhm.append(float(fwhm_val))
    return np.array(x), np.array(y), np.array(fwhm)

def estimate_dist(filename, mode, interpolated_x):
    """Estimate the distribution based on the given mode ('g' for Gaussian, 'l' for Lorentzian)."""
    x, y, fwhm = read_data(filename)
    y_min, y_max = y.min(), y.max()
    y_norm = (y - y_min) / (y_max - y_min)
    response_arr = []
    for x_val, y_val, fwhm_val in zip(x, y_norm, fwhm):
        if mode == 'l':
            distribution = lorentzian_fit(interpolated_x, x_val, y_val, fwhm_val)
        elif mode == 'g':
            distribution = gaussian_fit(interpolated_x, x_val, y_val, fwhm_val)
        else:
            raise ValueError(f'Unknown function type {mode}')
        response_arr.append(distribution)
    response_arr = np.array(response_arr)
    full_spectrum = response_arr.sum(axis=0)
    full_spectrum_max = full_spectrum.max()
    full_spectrum = full_spectrum / full_spectrum_max
    response_arr = response_arr / full_spectrum_max
    y_scaled = y_norm / full_spectrum_max
    return x, y_scaled, y_norm, response_arr, full_spectrum

def create_default_file(filename):
    """Create a default text file with specified values."""
    if "Experiment.txt" in filename:
        default_data = np.array([
            [19193.9, 0.33, 1500],
            [17543.9, 1.00, 1500],
            [16447.4, 0.93, 1500],
            [15822.8, 0.40, 1500],
            [14992.5, 0.64, 1500],
            [13624.0, 0.21, 1500],
            [10000.0, 0.00, 1500],
            [10000.0, 0.00, 1500],
            [10000.0, 0.00, 1500],
            [10000.0, 0.00, 1500]
        ])
        np.savetxt(filename, default_data, fmt=['%.1f', '%.2f', '%d'])


def main():
    """Main function to read configurations, estimate distributions, and plot the data."""
    current_dir = os.getcwd()  # Get current directory
    config = [
        {'filename': os.path.join(current_dir, 'Absorption_Input.txt'), 
         'mode': 'g', 
         'title':'Calculated', 
         'inv_axis':True, 
         'plot_full':True, 'plot_bar':True, 'plot_cont':True},
        {'filename': os.path.join(current_dir, 'Experiment.txt'), 
         'mode': mode, 
         'title':'Experiment ( or mock data)', 
         'inv_axis':True, 
         'plot_full':True, 'plot_bar':True, 'plot_cont':True}
    ]
    
    for conf in config:
        if not os.path.exists(conf['filename']):
            create_default_file(conf['filename'])
    
    figsize = (8,8)
    interpolated_x = np.linspace(x_start, x_end, int((x_end-x_start)/x_resolution))
    fig, axs = plt.subplots(nrows=len(config), ncols=1, sharex=False, sharey=False, figsize=figsize)
    axs = [axs] if len(config) <= 1 else axs
    for ax, conf in zip(axs, config):
        x, y_scaled, y_norm, response_arr, full_spectrum = estimate_dist(conf['filename'], conf['mode'], interpolated_x)
        if conf['plot_cont']:
            for i in range(np.size(y_scaled)): 
                ax.scatter(interpolated_x, response_arr[i, :], marker='.', alpha=0.5, linewidths=0.75)
        if conf['plot_full']:
            ax.plot(interpolated_x, full_spectrum, 'b-', linewidth='3')
        if conf['plot_bar']:
            ax.bar(x, y_scaled, color='red', width=0.025)
        ax.set_title(conf.get('title',''))
        if conf.get('inv_axis', True):
            ax_xmin, ax_xmax = ax.get_xlim()
            ax.set_xlim(ax_xmax, ax_xmin)
        output_filename = conf['filename'].split('/')[-1].split('.')[0] + '_output.csv'
        with open(output_filename, 'w') as outfile:
            header = 'interpolated_x,' + ','.join([f'peak_{i+1}' for i in range(response_arr.shape[0])]) + ',full_spectrum,,x,y_scaled,y_norm\n'
            outfile.write(header)
            for i in range(interpolated_x.shape[0]):
                row = str(interpolated_x[i])
                for j in range(response_arr.shape[0]):
                    row += ',' + str(response_arr[j, i])
                row += ',' + str(full_spectrum[i]) + ','
                if i < np.size(y_scaled):
                    row += ',' + str(x[i]) + ',' + str(y_scaled[i]) + ',' + str(y_norm[i])
                else:
                    row += ',,,'
                row += '\n'
                outfile.write(row)
    plt.show()

if __name__ == '__main__':
    main()
